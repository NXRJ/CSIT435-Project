"""Build the submission-ready CSCI435 project report from executed notebook artifacts."""

from __future__ import annotations

import json
import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS = ROOT / "artifacts"
OUTPUT_DIR = ROOT / "output" / "report"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = Path(os.environ.get("CSCI435_REPORT_OUTPUT", OUTPUT_DIR / "CSCI435_Project_Report.docx"))

METRICS = json.loads((ARTIFACTS / "metrics.json").read_text(encoding="utf-8"))
ROBUSTNESS = pd.read_csv(ARTIFACTS / "robustness_results.csv")
MODEL_COMPARISON = pd.read_csv(ARTIFACTS / "model_comparison.csv")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x17, 0x32, 0x4D)
MUTED = RGBColor(0x5C, 0x66, 0x70)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths_dxa, header=True, font_size=9.0, alignments=None):
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    if header:
        repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if row_index == 0 and header:
                shade_cell(cell, LIGHT_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if alignments:
                    paragraph.alignment = alignments[col_index]
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(row_index == 0 and header))


def set_alt_text(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_picture(doc, path, width, caption, alt_text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    set_alt_text(shape, alt_text)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    return shape


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_code_block(doc, code):
    paragraph = doc.add_paragraph(style="Code Block")
    for index, line in enumerate(code.strip().splitlines()):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "D9DEE5")
        borders.append(border)
    p_pr.append(borders)


def add_callout(doc, lead, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.16)
    run = paragraph.add_run(f"{lead} ")
    set_run_font(run, bold=True, color=DARK_BLUE)
    paragraph.add_run(text)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead:
        run = paragraph.add_run(bold_lead)
        set_run_font(run, bold=True)
    paragraph.add_run(text)
    return paragraph


doc = Document()
doc.core_properties.title = "Accessibility Scene Hazard Assistant - CSCI435 Project Report"
doc.core_properties.subject = "CSCI435 Computer Vision Algorithms and Systems"
doc.core_properties.author = "CSCI435 Project Group"
doc.core_properties.keywords = "computer vision, accessibility, OpenCV, SVM, Gradio"

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

heading_tokens = {
    "Heading 1": (16, BLUE, 16, 8),
    "Heading 2": (13, BLUE, 12, 6),
    "Heading 3": (12, DARK_BLUE, 8, 4),
}
for style_name, (size, color, before, after) in heading_tokens.items():
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True

for style_name in ("List Bullet", "List Number"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167

caption = styles["Caption"]
caption.font.name = "Calibri"
caption.font.size = Pt(9)
caption.font.italic = True
caption.font.color.rgb = MUTED

if "Code Block" not in styles:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code Block"]
code_style.font.name = "Consolas"
code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
code_style.font.size = Pt(8.5)
code_style.paragraph_format.left_indent = Inches(0.12)
code_style.paragraph_format.right_indent = Inches(0.12)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(7)
code_style.paragraph_format.line_spacing = 1.0
code_style.paragraph_format.keep_together = True

header = section.header
header_p = header.paragraphs[0]
header_p.text = "CSCI435  |  Accessibility Scene Hazard Assistant"
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
for run in header_p.runs:
    set_run_font(run, size=8.5, color=MUTED, bold=True)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_p.paragraph_format.space_before = Pt(0)
add_page_field(footer_p)
for run in footer_p.runs:
    set_run_font(run, size=8.5, color=MUTED)

# Editorial-cover pattern for a formal academic report.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(72)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(14)
set_run_font(kicker.add_run("UNIVERSITY OF WOLLONGONG IN DUBAI"), size=11, color=BLUE, bold=True)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(9)
set_run_font(title.add_run("Accessibility Scene\nHazard Assistant"), size=29, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(36)
set_run_font(subtitle.add_run("An Integrated Computer Vision Application for Accessible Scene Guidance"), size=14, color=DARK_BLUE)

course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course.paragraph_format.space_after = Pt(28)
set_run_font(course.add_run("CSCI435 - Computer Vision Algorithms and Systems\nSpring 2026"), size=11.5, color=NAVY, bold=True)

members = [
    ("Mehdi Leghmizi", "8528834"),
    ("Neeraj Santosh", "8329345"),
    ("Muhammad Soban", "8555588"),
    ("Zachary Bracke", "8947405"),
    ("Mostafa Shalash", "7391493"),
]
for name, student_id in members:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(f"{name}  |  {student_id}"), size=10.5, color=NAVY)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(28)
meta.paragraph_format.space_after = Pt(0)
set_run_font(meta.add_run("Lecturer: Dr Patrick Mukala\nReport date: 27 June 2026"), size=9.5, color=MUTED)

doc.add_page_break()

doc.add_heading("Report at a glance", level=1)
add_body(
    doc,
    "We developed a locally deployable web application that combines image enhancement, edge detection, "
    "object proposal, custom-trained object recognition, moving-object detection and object tracking in one "
    "workflow. The system accepts an uploaded image, webcam capture or video and returns annotated visual "
    "results, text guidance, latency and estimated throughput.",
)
add_callout(
    doc,
    "Key outcome.",
    f"The executed notebook measured {METRICS['mean_detection_precision']:.1%} mean precision, "
    f"{METRICS['mean_detection_recall']:.1%} recall and {METRICS['mean_detection_f1']:.1%} F1 across seven "
    f"scene conditions. The full still-frame pipeline averaged {METRICS['performance']['mean_fps']:.1f} FPS "
    f"({METRICS['performance']['mean_latency_ms']:.1f} ms), exceeding the 10 FPS requirement.",
)

doc.add_heading("Abstract", level=2)
add_body(
    doc,
    "This project investigates how several computer vision techniques can operate together as a usable system "
    "rather than as isolated laboratory exercises. The proof-of-concept is an accessibility-oriented scene "
    "assistant for recognising three marker types: stop, warning and safe route. A deterministic custom dataset "
    "of 560 augmented marker crops is generated inside the notebook. HOG, HSV, intensity and Hu-moment features "
    "are used to compare an RBF support vector machine, random forest and k-nearest-neighbour classifier. The "
    "selected SVM is then integrated with colour/contour proposals, adaptive Canny edges, CLAHE enhancement, "
    "MOG2 background modelling and centroid tracking. Evaluation covers held-out recognition, class-aware scene "
    "detection, lighting and orientation changes, noise, blur, occlusion, clutter, blank inputs and CPU speed. "
    "Results are strong within the controlled synthetic domain, but the report also identifies the synthetic-to-real "
    "gap and the need for user testing before any real assistive deployment.",
)

doc.add_heading("Submission components", level=2)
for item in (
    "Self-contained and executed Jupyter notebook: CSCI435_Project.ipynb.",
    "Dependency file and reproducible execution instructions in README.md.",
    "Two-minute-five-second demonstration video generated from actual pipeline outputs.",
    "Editable DOCX report and submission-ready PDF report.",
    "Timed defence guide, robust sample inputs and requirement-to-evidence matrix.",
):
    add_bullet(doc, item)

doc.add_page_break()
doc.add_heading("Introduction", level=1)
doc.add_heading("Problem statement", level=2)
add_body(
    doc,
    "People with limited vision may need a quick indication of warnings, safe-route cues, visible boundaries and "
    "movement in a camera view. Many classroom demonstrations process one image with one algorithm, which does "
    "not show how a usable application should validate input, combine results or communicate uncertainty. Our goal "
    "was therefore to create one deployable interface that links low-level, mid-level and high-level vision tasks "
    "around a clear accessibility user story.",
)

doc.add_heading("User story", level=2)
add_callout(
    doc,
    "User story.",
    "As a user, I want to upload an image or video, or capture a webcam image, so that one interface enhances the "
    "scene, recognises marker-like objects, highlights boundaries and movement, and gives concise guidance.",
)

doc.add_heading("Selected vision capabilities", level=2)
capabilities = [
    ("Image enhancement", "gray-world colour balancing followed by CLAHE on the luminance channel."),
    ("Edge detection", "adaptive Canny thresholds with morphological closing and a visible edge overlay."),
    ("Object detection", "HSV colour segmentation, morphology, contour filtering and bounding boxes."),
    ("Object recognition", "an RBF SVM trained on the group's generated custom marker dataset."),
    ("Change and moving-object detection", "MOG2 adaptive background modelling on video frames."),
    ("Object tracking", "centroid association with persistent IDs and missing-frame tolerance."),
]
for label, detail in capabilities:
    add_bullet(doc, f"{label}: {detail}")

doc.add_heading("Scope and safety", level=2)
add_body(
    doc,
    "The application is an academic proof-of-concept, not a certified safety device. Its marker vocabulary and "
    "synthetic evaluation are deliberately constrained so that every training and testing step is reproducible. "
    "Real-world assistive use would require representative data, calibrated depth, accessibility studies, privacy "
    "review and a much more conservative false-negative analysis.",
)

doc.add_heading("System architecture", level=1)
add_picture(
    doc,
    ARTIFACTS / "architecture_diagram.png",
    6.35,
    "Figure 1. System architecture and data flow.",
    "Architecture diagram showing image, webcam or video input, validation, enhancement, edges, object proposals, SVM recognition, motion tracking, fusion and Gradio output.",
)
add_body(
    doc,
    "The frontend and backend are packaged in the notebook. Gradio provides the browser interface while Python, "
    "OpenCV and scikit-learn provide the processing and model layers. An input is validated and resized to a maximum "
    "width of 640 pixels. The enhanced frame feeds edge and colour-proposal branches. Candidate pixels are "
    "canonicalised and classified using the custom SVM. Video frames also pass through MOG2 and the centroid "
    "tracker. Fusion draws boxes, labels and track IDs, builds a text guidance sentence and reports latency/FPS.",
)

doc.add_heading("Frontend and input modalities", level=2)
add_body(
    doc,
    "The Image / Webcam tab accepts an upload or camera capture and shows four outputs: the integrated annotated "
    "result, enhanced image, edge view and guidance. The Video tab accepts an uploaded file, applies the same frame "
    "pipeline plus motion/tracking, and returns an annotated video and aggregate metrics. These two paths satisfy "
    "the requirement for at least two input modalities while keeping the interface consistent.",
)
add_picture(
    doc,
    ARTIFACTS / "gradio_interface.png",
    5.7,
    "Figure 2. Verified local Gradio interface in the Image / Webcam mode.",
    "Screenshot of the running Gradio application with input, integrated result, guidance, enhanced view, edge view and reliable defence example.",
)

doc.add_heading("Implementation details", level=1)
doc.add_heading("Enhancement and edge processing", level=2)
add_body(
    doc,
    "Gray-world balancing estimates a target mean across BGR channels and limits channel scaling to avoid extreme "
    "colour shifts. The result is converted to CIELAB, where CLAHE is applied only to luminance. Adaptive Canny "
    "thresholds are calculated from the median of a Gaussian-smoothed grayscale image; a 3 x 3 closing operation "
    "joins short breaks. This branch supplies a clear boundary view but does not directly claim semantic meaning.",
)

doc.add_heading("Object proposal and custom recognition", level=2)
add_body(
    doc,
    "HSV thresholds form red, yellow and green proposal masks. Opening removes isolated pixels and closing joins "
    "small gaps. Original and enhanced pixels are both searched using separate colour masks so a dominant single "
    "colour cannot be lost during balancing and adjacent marker colours cannot merge. Contours below 300 pixels or outside broad aspect-ratio limits are rejected. Each remaining object "
    "is centred on a neutral 96 x 96 canvas at a fixed foreground scale. This canonicalisation solved an important "
    "training/inference mismatch: the recogniser now sees the same object scale during both phases. Proposal masks "
    "are classified by the SVM using pre-enhancement validated pixels to preserve the training appearance distribution. "
    "When the synthetic crop model is uncertain, a conservative fusion rule checks canonical evidence: red octagonal "
    "regions with a bright symbol support stop, yellow triangular regions support warning, and green circular regions "
    "with a bright internal symbol support safe. This recovered real-style examples without replacing the learned model.",
)
add_code_block(
    doc,
    """
enhanced = enhance_image(frame)
edges = adaptive_canny(enhanced)
detections = detect_markers(
    enhanced, self.model, classification_frame=frame
)
""",
)

doc.add_heading("Motion detection and tracking", level=2)
add_body(
    doc,
    "For video, OpenCV's MOG2 model learns an adaptive background. Shadow-valued pixels are removed by a high "
    "binary threshold, followed by opening and dilation. After four warm-up frames, contours between 500 pixels "
    "and 65% of the frame area become moving regions. A lightweight centroid tracker associates each current box "
    "with the nearest existing centre, assigns a new ID when required and removes IDs after eight missed frames. "
    "This method is fast and easy to explain, although it can exchange identities when objects cross.",
)

doc.add_heading("Input validation and robustness controls", level=2)
for item in (
    "Rejects null, empty, malformed and smaller-than-32-pixel inputs with readable errors.",
    "Handles grayscale and four-channel images by converting them to three-channel BGR.",
    "Resizes wide inputs while preserving aspect ratio to bound latency and memory use.",
    "Returns a clear no-marker message for blank or unmatched scenes instead of failing.",
    "Uses MP4 output with an AVI fallback if the local codec cannot initialise.",
):
    add_bullet(doc, item)

doc.add_heading("Custom data and model training", level=1)
doc.add_heading("Dataset construction", level=2)
add_body(
    doc,
    f"The notebook generates {METRICS['dataset']['total_images']} labelled images, with "
    f"{METRICS['dataset']['per_class']} examples in each of four classes: stop, warning, safe and other. Variations "
    "include background level and gradient, scale, rotation, brightness, blur, Gaussian noise and partial occlusion. "
    "The other class is a red square with a diagonal line, intentionally designed to pass the colour proposal stage "
    "and force the recogniser to use shape evidence. Random seeds are fixed for reproducibility.",
)
add_picture(
    doc,
    ARTIFACTS / "dataset_montage.png",
    5.25,
    "Figure 3. Examples from the custom synthetic dataset.",
    "Montage of stop, warning, safe and hard-negative marker training images under varied conditions.",
)

doc.add_heading("Features and candidate models", level=2)
add_body(
    doc,
    "Each canonical crop is represented by HOG gradients, 18-bin hue and 12-bin saturation histograms, four "
    "intensity/colour statistics and seven log-transformed Hu moments. We used a stratified 75/25 split and compared "
    "three lightweight classifiers on the same samples. Scaling is included inside the SVM and k-NN pipelines to "
    "avoid information leakage from the held-out split.",
)

model_table = doc.add_table(rows=1, cols=4)
for index, text in enumerate(("Model", "Validation accuracy", "Training (s)", "Inference (ms/crop)")):
    model_table.rows[0].cells[index].text = text
for _, row in MODEL_COMPARISON.iterrows():
    cells = model_table.add_row().cells
    cells[0].text = str(row["model"])
    cells[1].text = f"{row['validation_accuracy']:.1%}"
    cells[2].text = f"{row['training_seconds']:.3f}"
    cells[3].text = f"{row['inference_ms_per_crop']:.3f}"
style_table(
    model_table,
    [3000, 2100, 1800, 2460],
    font_size=8.7,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
)
source = doc.add_paragraph("Table 1. Candidate-model comparison from the executed notebook.")
source.style = "Caption"
source.alignment = WD_ALIGN_PARAGRAPH.CENTER
source.paragraph_format.space_before = Pt(4)
source.paragraph_format.space_after = Pt(8)

doc.add_heading("Model choice", level=2)
add_body(
    doc,
    "All three models separated the controlled crop dataset perfectly, so crop accuracy alone could not justify a "
    "choice. We selected the RBF SVM because it combined a non-linear boundary, class balancing, probability output "
    "for user-facing confidence and low per-crop inference time. The trained pipeline is saved as a joblib artifact "
    "and is also recreated from source whenever the notebook is run.",
)
add_picture(
    doc,
    ARTIFACTS / "confusion_matrix.png",
    4.65,
    "Figure 4. Held-out confusion matrix for the selected SVM.",
    "Four-by-four confusion matrix showing 35 correct predictions in each synthetic held-out class and no errors.",
)
add_callout(
    doc,
    "Interpretation warning.",
    "The 100% held-out crop accuracy applies only to the generated marker distribution. It must not be presented as "
    "real-world safety accuracy. End-to-end scene metrics are more demanding and are reported separately.",
)

doc.add_heading("Experiments and results", level=1)
doc.add_heading("Evaluation protocol", level=2)
add_body(
    doc,
    "Recognition robustness used 35 unseen crops per class and condition. End-to-end detection used 12 scenes per "
    "condition, with one to three target classes per scene and seeds not used for training. A prediction counted as "
    "a true positive only when its class matched and its intersection-over-union with an unused ground-truth box "
    "was at least 0.30. Precision, recall, F1 and mean matched IoU were then aggregated. The seven scene conditions "
    "were bright, dim, rotated, blurred, noisy, occluded and cluttered.",
)

doc.add_page_break()
robust_table = doc.add_table(rows=1, cols=6)
for index, text in enumerate(("Condition", "Precision", "Recall", "F1", "Mean IoU", "Recognition accuracy")):
    robust_table.rows[0].cells[index].text = text
for _, row in ROBUSTNESS.iterrows():
    cells = robust_table.add_row().cells
    cells[0].text = str(row["condition"]).title()
    cells[1].text = f"{row['precision']:.1%}"
    cells[2].text = f"{row['recall']:.1%}"
    cells[3].text = f"{row['f1']:.1%}"
    cells[4].text = f"{row['mean_iou']:.3f}"
    cells[5].text = "N/A" if pd.isna(row["recognition_accuracy"]) else f"{row['recognition_accuracy']:.1%}"
style_table(
    robust_table,
    [1800, 1320, 1200, 1080, 1320, 2640],
    font_size=8.1,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 5,
)
source = doc.add_paragraph("Table 2. Condition-wise robustness results from the executed notebook.")
source.style = "Caption"
source.alignment = WD_ALIGN_PARAGRAPH.CENTER
source.paragraph_format.space_before = Pt(4)
source.paragraph_format.space_after = Pt(8)

add_body(
    doc,
    f"Across the seven conditions, mean precision was {METRICS['mean_detection_precision']:.1%}, mean recall was "
    f"{METRICS['mean_detection_recall']:.1%} and mean F1 was {METRICS['mean_detection_f1']:.1%}. Blurred and noisy "
    "generated scenes reached perfect detection scores in this run. Rotated, occluded and cluttered scenes were "
    "harder. Clutter reduced precision because saturated distractors sometimes survived proposal filtering; rotated "
    "and occluded targets mainly reduced recall. This pattern is consistent with the proposal/classification design.",
)
add_picture(
    doc,
    ARTIFACTS / "robustness_chart.png",
    6.0,
    "Figure 5. Detection and recognition robustness by condition.",
    "Bar chart comparing detection precision, detection recall and recognition accuracy across seven conditions.",
)

doc.add_heading("Performance", level=2)
performance_rows = [
    ("Still-frame mean latency", f"{METRICS['performance']['mean_latency_ms']:.2f} ms"),
    ("Still-frame median latency", f"{METRICS['performance']['median_latency_ms']:.2f} ms"),
    ("Still-frame 95th-percentile latency", f"{METRICS['performance']['p95_latency_ms']:.2f} ms"),
    ("Estimated still-frame throughput", f"{METRICS['performance']['mean_fps']:.2f} FPS"),
    ("Processed-video mean latency", f"{METRICS['video']['mean_latency_ms']:.2f} ms/frame"),
    ("Processed-video estimated throughput", f"{METRICS['video']['mean_fps']:.2f} FPS"),
    ("Rubric threshold", "10 FPS - met"),
]
perf_table = doc.add_table(rows=1, cols=2)
perf_table.rows[0].cells[0].text = "Measure"
perf_table.rows[0].cells[1].text = "Result"
for label, value in performance_rows:
    cells = perf_table.add_row().cells
    cells[0].text = label
    cells[1].text = value
style_table(
    perf_table,
    [5700, 3660],
    font_size=9.0,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
)
source = doc.add_paragraph("Table 3. CPU timing measured over the complete integrated pipeline.")
source.style = "Caption"
source.alignment = WD_ALIGN_PARAGRAPH.CENTER
source.paragraph_format.space_before = Pt(4)
source.paragraph_format.space_after = Pt(8)
add_body(
    doc,
    "The benchmark includes validation, resizing, enhancement, edges, proposal generation, feature extraction, "
    "SVM inference, drawing and guidance. Video timing also includes MOG2 and tracking. FPS is computed as the "
    "reciprocal of processing latency and therefore measures processing capacity rather than camera-display refresh "
    "rate. Results depend on hardware, but the tested run had a substantial margin above the 10 FPS criterion.",
)

doc.add_heading("Qualitative results", level=2)
add_picture(
    doc,
    ARTIFACTS / "qualitative_pipeline.png",
    6.15,
    "Figure 6. End-to-end image workflow and generated guidance.",
    "Four-panel qualitative output showing the original scene, enhanced scene, Canny overlay and detected stop, warning and safe markers.",
)
add_body(
    doc,
    "The clean defence scene demonstrates how the tasks support one another. Enhancement normalises contrast, the "
    "edge view exposes structural boundaries, colour/contour proposals localise candidates, the SVM labels all three "
    "markers and the fusion layer produces: 'Safe-route marker on the right; warning marker in the centre; stop marker "
    "on the left.' Video examples add motion boxes and persistent IDs using the same base pipeline.",
)

doc.add_heading("Correctness and edge-case tests", level=2)
for item in (
    "Blank uniform scene: zero marker detections and a readable no-marker message.",
    "Tiny 12 x 12 image: rejected with a ValueError instead of entering OpenCV operations.",
    "Dim, rotated, blurred, noisy and occluded scenes: processed and scored condition by condition.",
    "Cluttered scene: false positives are measured rather than removed from the reported results.",
    f"Supplied real-style regression images: stop detected at {METRICS['real_style_regression']['stop']['confidence']:.1%} and safe detected at {METRICS['real_style_regression']['safe']['confidence']:.1%}.",
    "Generated 90-frame video: successfully encoded, processed and measured with motion/tracking outputs.",
):
    add_bullet(doc, item)

doc.add_heading("Challenges and critical analysis", level=1)
doc.add_heading("Training-to-inference mismatch", level=2)
add_body(
    doc,
    "An early version classified padded scene crops directly even though training used tightly framed 96 x 96 "
    "examples. Crop accuracy looked perfect, but scene recall was only 7.1%. This was a useful warning against "
    "reporting a convenient metric without testing the complete system. Canonical foreground scaling and recognition "
    "from pre-enhancement pixels raised the verified mean scene F1 to the value reported in Table 2.",
)

doc.add_heading("Remaining limitations", level=2)
limitations = [
    "Synthetic-to-real domain gap: natural signs, reflections, weathering and camera compression are not represented.",
    "Colour dependence: extreme colour casts or desaturation may prevent proposal generation before the SVM runs.",
    "No calibrated distance: left/centre/right guidance is relative image position, not metric depth or collision risk.",
    "Motion is not semantic: MOG2 identifies changing regions but cannot state whether they are people or vehicles.",
    "Simple tracking: nearest-centroid matching may exchange IDs during crossing or prolonged occlusion.",
    "Safety and accessibility: no target-user study, audio/haptic interface or formal false-negative safety threshold was completed.",
]
for item in limitations:
    add_bullet(doc, item)

doc.add_heading("UI and deployment trade-offs", level=2)
add_body(
    doc,
    "Gradio made the application fast to reproduce and suitable for a live local defence. It also exposes upload and "
    "webcam sources without maintaining a separate JavaScript codebase. The trade-off is less control over mobile "
    "layout and browser codec behaviour than a custom frontend. Local hosting avoids sending camera data to a cloud "
    "service, but the prototype does not yet implement persistent privacy controls or authentication.",
)

doc.add_heading("Individual contributions", level=1)
add_body(
    doc,
    "The following table records the agreed primary workstreams. Integration, review and defence preparation were "
    "shared so that each member contributes to coding, documentation and presentation as required.",
)
contrib_table = doc.add_table(rows=1, cols=5)
for index, text in enumerate(("Member", "Student ID", "Primary coding", "Documentation", "Defence role")):
    contrib_table.rows[0].cells[index].text = text
contributions = [
    ("Mehdi Leghmizi", "8528834", "UI integration and input handling", "User story and requirements", "Opening and UI demo"),
    ("Neeraj Santosh", "8329345", "Custom dataset and model training", "Training methodology", "Model selection"),
    ("Muhammad Soban", "8555588", "Enhancement, edges and robustness tests", "Experiment methodology", "Image robustness"),
    ("Zachary Bracke", "8947405", "Motion, tracking and performance", "Results and critical analysis", "Video and performance"),
    ("Mostafa Shalash", "7391493", "System integration and QA", "Architecture, conclusion and references", "Architecture and closing"),
]
for row in contributions:
    cells = contrib_table.add_row().cells
    for index, value in enumerate(row):
        cells[index].text = value
style_table(
    contrib_table,
    [1900, 1180, 2360, 2100, 1820],
    font_size=7.8,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)
source = doc.add_paragraph("Table 4. Group contribution record and defence responsibilities.")
source.style = "Caption"
source.alignment = WD_ALIGN_PARAGRAPH.CENTER
source.paragraph_format.space_before = Pt(4)
source.paragraph_format.space_after = Pt(8)

doc.add_heading("Conclusion and future work", level=1)
add_body(
    doc,
    "The project demonstrates a working integrated system rather than a collection of isolated scripts. Six "
    "permissible computer vision capabilities operate inside one image/video workflow and are exposed through a "
    "responsive local web interface. The custom SVM satisfies the training requirement, the end-to-end evaluation "
    "tests robustness rather than relying only on crop accuracy, and the measured performance is comfortably above "
    "the real-time threshold. The most important lesson was that complete-system evaluation can reveal failures hidden "
    "by excellent component metrics.",
)
add_body(
    doc,
    "Future work should collect consented real images and videos, annotate objects with bounding boxes, compare a "
    "fine-tuned detector with the current proposal-plus-classifier approach, calibrate depth, quantify false negatives "
    "by hazard severity and test audio/haptic guidance with target users. A mobile deployment would also need model "
    "compression, battery profiling, offline inference and privacy-by-design review.",
)

doc.add_heading("References", level=1)
references = [
    "Bradski, G. (2000). The OpenCV Library. Dr. Dobb's Journal of Software Tools.",
    "Canny, J. (1986). A computational approach to edge detection. IEEE Transactions on Pattern Analysis and Machine Intelligence, 8(6), 679-698. https://doi.org/10.1109/TPAMI.1986.4767851",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20, 273-297. https://doi.org/10.1007/BF00994018",
    "Dalal, N., & Triggs, B. (2005). Histograms of oriented gradients for human detection. Proceedings of CVPR, 886-893. https://doi.org/10.1109/CVPR.2005.177",
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
    "Zivkovic, Z. (2004). Improved adaptive Gaussian mixture model for background subtraction. Proceedings of ICPR, 28-31. https://doi.org/10.1109/ICPR.2004.1333992",
    "OpenCV. (2026). OpenCV documentation. https://docs.opencv.org/4.x/",
    "Scikit-learn developers. (2026). SVC documentation. https://scikit-learn.org/stable/modules/generated/sklearn.svm.SVC.html",
    "Gradio. (2026). Gradio documentation. https://www.gradio.app/docs",
    "University of Wollongong in Dubai. (2026). CSCI435 project description, Spring 2026.",
]
for reference in references:
    paragraph = doc.add_paragraph(reference)
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.first_line_indent = Inches(-0.3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True

doc.add_page_break()
doc.add_heading("Appendix A - Rubric traceability", level=1)
rubric_rows = [
    ("Four or more tasks", "15", "Six integrated capabilities in VisionSystem; Figures 1 and 6."),
    ("Correctness and robustness", "15", "Seven conditions, edge-case assertions and class-aware IoU evaluation; Table 2."),
    ("UI and UX", "10", "Verified Gradio image/webcam and video tabs; Figure 2."),
    ("Model selection and training", "10", "560 custom images, three-model comparison and saved SVM; Table 1."),
    ("Architecture and code quality", "10", "Typed functions/classes, unified notebook, README and architecture diagram."),
    ("Performance", "5", f"{METRICS['performance']['mean_fps']:.1f} FPS measured versus 10 FPS requirement; Table 3."),
    ("Report quality", "15", "Required sections, diagrams, quantitative results, qualitative examples and critical analysis."),
    ("Live demonstration", "15", "Local interface, stable generated examples, 2:05 backup video and defence guide."),
    ("Teamwork", "5", "Named coding, documentation and defence responsibilities; Table 4."),
]
rubric_table = doc.add_table(rows=1, cols=3)
for index, text in enumerate(("Criterion", "Marks", "Evidence")):
    rubric_table.rows[0].cells[index].text = text
for row in rubric_rows:
    cells = rubric_table.add_row().cells
    for index, value in enumerate(row):
        cells[index].text = value
style_table(
    rubric_table,
    [3100, 900, 5360],
    font_size=8.2,
    alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

doc.add_heading("Appendix B - Reproduction and defence checklist", level=1)
checklist = [
    "Create a Python 3.12 virtual environment and install requirements.txt.",
    "Open CSCI435_Project.ipynb and run all cells from top to bottom.",
    "Confirm that the model comparison, confusion matrix, robustness table and performance assertion appear.",
    "Open the printed local Gradio URL and process the generated defence image and sample video.",
    "Show image enhancement, edge overlay, marker boxes, guidance, motion IDs and measured FPS.",
    "Keep artifacts/CSCI435_Demonstration_Video.mp4 available as the required video and technical-failure backup.",
    "Push the repository to GitHub or GitLab, grant the lecturer access and place the repository link in the Moodle submission.",
]
for item in checklist:
    add_bullet(doc, item)

update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
doc.settings.element.append(update_fields)
doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
